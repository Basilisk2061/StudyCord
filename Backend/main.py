import os
import time
import hashlib
# pyrefly: ignore [missing-import]
import httpx
# pyrefly: ignore [missing-import]
from fastapi import FastAPI, HTTPException
# pyrefly: ignore [missing-import]
from fastapi.middleware.cors import CORSMiddleware
# pyrefly: ignore [missing-import]
from dotenv import load_dotenv
from pathlib import Path

# Load env variables from Backend/.env using absolute path
BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")

import uuid
import json
import re
import io
import secrets
import string
from datetime import datetime, timezone
from fastapi import UploadFile, File, Header, Depends
from pydantic import BaseModel

# LangChain and Gemini imports (embeddings only)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# OpenRouter LLM import
from langchain_openai import ChatOpenAI

# Map GEMINI_API_KEY to GOOGLE_API_KEY for langchain-google-genai compatibility
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY and not os.getenv("GOOGLE_API_KEY"):
    os.environ["GOOGLE_API_KEY"] = GEMINI_API_KEY

# Load OpenRouter configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "openrouter/auto")

# =====================================================================
# STARTUP LOGGING — safe, never prints secrets
# =====================================================================
print("=" * 60)
print("[BACKEND] StudyCord RAG Backend starting up...")
if os.getenv("GOOGLE_API_KEY"):
    print("[BACKEND] [OK] GEMINI_API_KEY/GOOGLE_API_KEY loaded (used for embeddings).")
else:
    print("[BACKEND] [!!] WARNING: GEMINI_API_KEY is not configured -- embeddings will fail.")

if OPENROUTER_API_KEY:
    masked = f"{OPENROUTER_API_KEY[:8]}...{OPENROUTER_API_KEY[-4:]}"
    print(f"[BACKEND] [OK] OPENROUTER_API_KEY loaded ({masked}).")
else:
    print("[BACKEND] [!!] WARNING: OPENROUTER_API_KEY is not configured -- generation will fail.")

print(f"[BACKEND] OpenRouter model: {OPENROUTER_MODEL}")
print("=" * 60)

METERED_DOMAIN = os.getenv("METERED_DOMAIN", "studycord.metered.live")
METERED_SECRET_KEY = os.getenv("METERED_SECRET_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")

app = FastAPI(title="StudyCord Secure TURN API")

# Enable CORS for the frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:5174",
        "http://127.0.0.1:5174",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =====================================================================
# OPENROUTER HELPER
# =====================================================================

def get_rag_chat_model(temperature: float = 0.3) -> ChatOpenAI:
    """
    Returns a ChatOpenAI instance routed through OpenRouter.
    Temperature can be tuned per-task (lower for structured, higher for chat).
    """
    if not OPENROUTER_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="OPENROUTER_API_KEY is not configured on the backend. Please add it to Backend/.env"
        )

    return ChatOpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=OPENROUTER_API_KEY,
        model=OPENROUTER_MODEL,
        temperature=temperature,
        request_timeout=120,
        default_headers={
            "HTTP-Referer": "http://localhost:5173",
            "X-Title": "StudyCord",
        },
    )


def _handle_openrouter_error(e: Exception, endpoint: str) -> HTTPException:
    """
    Translates OpenRouter / upstream errors into friendly frontend messages.
    Returns an HTTPException ready to be raised.
    """
    error_str = str(e).lower()
    detail_raw = str(e)

    # Try to extract status code from the error
    status_code = 500

    if "401" in error_str or "invalid api key" in error_str or "unauthorized" in error_str:
        status_code = 401
        detail = "OpenRouter API key is invalid or expired. Please check your OPENROUTER_API_KEY in Backend/.env."
    elif "402" in error_str or "insufficient" in error_str or "payment" in error_str or "credits" in error_str:
        status_code = 402
        detail = "OpenRouter account has insufficient credits. Please add credits or switch to a free model."
    elif "429" in error_str or "rate limit" in error_str or "too many requests" in error_str:
        status_code = 429
        detail = "Rate limit reached on the AI model. Please wait a moment and try again."
    elif "timeout" in error_str or "timed out" in error_str:
        status_code = 504
        detail = "The AI model took too long to respond. Please try again."
    elif "model" in error_str and ("not found" in error_str or "unavailable" in error_str or "not available" in error_str):
        status_code = 503
        detail = f"The selected AI model ({OPENROUTER_MODEL}) is currently unavailable. Please try again later or change the model in Backend/.env."
    elif "content" in error_str and ("filter" in error_str or "safety" in error_str or "moderation" in error_str):
        status_code = 422
        detail = "The AI model refused to generate a response for this content. Please try rephrasing your request."
    else:
        detail = f"AI generation failed. Please try again. (Error: {detail_raw[:200]})"

    print(f"[{endpoint}] OpenRouter error — status={status_code}, detail={detail}")
    return HTTPException(status_code=status_code, detail=detail)


# =====================================================================
# SERVER ADMINISTRATION HELPERS
# =====================================================================

PERMISSIONS = {
    "owner": {
        "view_server",
        "manage_server",
        "manage_channels",
        "manage_members",
        "manage_roles",
        "kick_members",
        "ban_members",
        "create_invites",
        "manage_invites",
        "transfer_ownership",
        "delete_server",
    },
    "admin": {
        "view_server",
        "manage_server",
        "manage_channels",
        "manage_members",
        "kick_members",
        "ban_members",
        "create_invites",
        "manage_invites",
    },
    "member": {"view_server"},
}

VALID_ROLES = {"owner", "admin", "member"}
ROLE_RANK = {"member": 1, "admin": 2, "owner": 3}
DEFAULT_CHANNELS = ("general", "assignments", "resources")


class RoleUpdateRequest(BaseModel):
    role: str


class TransferOwnershipRequest(BaseModel):
    new_owner_id: str


class ReasonRequest(BaseModel):
    reason: str | None = None


class ServerUpdateRequest(BaseModel):
    name: str | None = None
    description: str | None = None


class CreateServerRequest(BaseModel):
    name: str


class JoinServerRequest(BaseModel):
    invite_code: str


class CreateChannelRequest(BaseModel):
    name: str
    type: str = "text"


def _supabase_configured() -> bool:
    return bool(SUPABASE_URL and SUPABASE_ANON_KEY)


class SupabaseRestClient:
    """Small PostgREST client with one fixed security context."""

    def __init__(self, name: str, api_key: str, bearer_token: str):
        self.name = name
        self.api_key = api_key
        self.bearer_token = bearer_token

    def _headers(self, prefer: str | None = None) -> dict[str, str]:
        headers = {
            "apikey": self.api_key,
            "Authorization": f"Bearer {self.bearer_token}",
            "Content-Type": "application/json",
        }
        if prefer:
            headers["Prefer"] = prefer
        return headers

    async def rest(self, method: str, path: str, *, params: dict | None = None, json_body=None, prefer: str | None = None):
        url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
        async with httpx.AsyncClient(timeout=20.0) as client:
            response = await client.request(
                method,
                url,
                params=params,
                json=json_body,
                headers=self._headers(prefer),
            )
        if response.status_code >= 400:
            print(f"[SUPABASE:{self.name}] {method} {path} failed: {response.status_code} {response.text[:400]}")
            status_code = response.status_code if response.status_code in {400, 401, 403, 404, 409} else 500
            raise HTTPException(status_code=status_code, detail="Database operation was rejected.")
        return response.json() if response.text else None

    async def rpc(self, function_name: str, payload: dict):
        url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/rpc/{function_name}"
        async with httpx.AsyncClient(timeout=20.0) as client:
            response = await client.post(url, json=payload, headers=self._headers())
        if response.status_code >= 400:
            print(f"[SUPABASE-RPC:{self.name}] {function_name} failed: {response.status_code} {response.text[:400]}")
            if "must be server owner" in response.text:
                raise HTTPException(status_code=403, detail="Only the server owner can transfer ownership.")
            if "new owner must be a current server member" in response.text:
                raise HTTPException(status_code=400, detail="New owner must be a current server member.")
            raise HTTPException(status_code=500, detail="Database function failed. Ensure the Phase 14 migration has been run.")
        return response.json() if response.text else None


def supabase_user(access_token: str) -> SupabaseRestClient:
    if not _supabase_configured():
        raise HTTPException(
            status_code=500,
            detail="Add SUPABASE_URL and SUPABASE_ANON_KEY to Backend/.env.",
        )
    return SupabaseRestClient("user", SUPABASE_ANON_KEY, access_token)


def supabase_admin() -> SupabaseRestClient:
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(
            status_code=500,
            detail="This operation requires SUPABASE_SERVICE_ROLE_KEY in Backend/.env.",
        )
    return SupabaseRestClient("admin", SUPABASE_SERVICE_ROLE_KEY, SUPABASE_SERVICE_ROLE_KEY)


async def get_current_user(authorization: str | None = Header(default=None)):
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing authentication token.")
    token = authorization.split(" ", 1)[1].strip()
    if not token:
        raise HTTPException(status_code=401, detail="Missing authentication token.")

    if not _supabase_configured():
        raise HTTPException(status_code=500, detail="Supabase auth credentials are not configured on the backend.")

    async with httpx.AsyncClient(timeout=10.0) as client:
        response = await client.get(
            f"{SUPABASE_URL.rstrip('/')}/auth/v1/user",
            headers={"apikey": SUPABASE_ANON_KEY, "Authorization": f"Bearer {token}"},
        )
    if response.status_code >= 400:
        raise HTTPException(status_code=401, detail="Invalid or expired authentication token.")
    user_data = response.json()
    user_id = user_data.get("id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid authentication token.")
    return {"id": user_id, "email": user_data.get("email"), "supabase_user": supabase_user(token)}


def _audit(action: str, server_id: str, actor_id: str, target_id: str | None = None, old_role: str | None = None, new_role: str | None = None, success: bool = True):
    print(json.dumps({
        "action": action,
        "server_id": server_id,
        "actor_user_id": actor_id,
        "target_user_id": target_id,
        "old_role": old_role,
        "new_role": new_role,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "success": success,
    }))


def generate_invite_code(length: int = 6) -> str:
    chars = string.ascii_uppercase + string.digits
    return "".join(secrets.choice(chars) for _ in range(length))


def normalize_channel_name(name: str) -> str:
    clean = re.sub(r"[^a-zA-Z0-9\s_-]", "", name.strip().lower())
    clean = re.sub(r"\s+", "-", clean)
    clean = re.sub(r"-+", "-", clean).strip("-_")
    if not clean:
        raise HTTPException(status_code=400, detail="Channel name is required.")
    return clean[:80]


async def get_server(client: SupabaseRestClient, server_id: str):
    rows = await client.rest("GET", "servers", params={"id": f"eq.{server_id}", "select": "*"})
    if not rows:
        raise HTTPException(status_code=404, detail="Server not found.")
    return rows[0]


async def get_server_member(client: SupabaseRestClient, server_id: str, user_id: str):
    rows = await client.rest(
        "GET",
        "server_members",
        params={"server_id": f"eq.{server_id}", "user_id": f"eq.{user_id}", "select": "*"},
    )
    return rows[0] if rows else None


async def get_server_member_role(client: SupabaseRestClient, server_id: str, user_id: str) -> str | None:
    member = await get_server_member(client, server_id, user_id)
    return member.get("role") if member else None


async def require_server_permission(client: SupabaseRestClient, server_id: str, user_id: str, permission: str) -> str:
    role = await get_server_member_role(client, server_id, user_id)
    if not role or permission not in PERMISSIONS.get(role, set()):
        raise HTTPException(status_code=403, detail="You do not have permission to perform this action.")
    return role


def can_manage_target_member(actor_role: str, target_role: str, action: str) -> bool:
    if target_role == "owner":
        return False
    if action in {"kick", "ban"}:
        return ROLE_RANK[actor_role] > ROLE_RANK[target_role]
    return actor_role == "owner"


def validate_role_change(actor_id: str, target_id: str, actor_role: str, target_role: str, new_role: str):
    if new_role not in {"admin", "member"}:
        raise HTTPException(status_code=400, detail="Role must be admin or member.")
    if actor_id == target_id:
        raise HTTPException(status_code=400, detail="You cannot change your own role.")
    if actor_role != "owner":
        raise HTTPException(status_code=403, detail="Only the server owner can change member roles.")
    if target_role == "owner":
        raise HTTPException(status_code=400, detail="The server owner role cannot be changed here.")
    if target_role == new_role:
        raise HTTPException(status_code=400, detail=f"Member is already {new_role}.")


async def cleanup_voice_presence(client: SupabaseRestClient, server_id: str, user_id: str):
    await client.rest(
        "DELETE",
        "voice_participants",
        params={"server_id": f"eq.{server_id}", "user_id": f"eq.{user_id}"},
    )


async def ensure_profile(client: SupabaseRestClient, user: dict):
    existing = await client.rest("GET", "profiles", params={"id": f"eq.{user['id']}", "select": "id"})
    if existing:
        return
    email = user.get("email") or ""
    username = email.split("@")[0] if email else f"user-{user['id'][:8]}"
    await client.rest(
        "POST",
        "profiles",
        json_body={"id": user["id"], "email": email, "username": username},
        prefer="return=minimal",
    )


@app.post("/api/servers")
async def create_server(request: CreateServerRequest, user=Depends(get_current_user)):
    client = user["supabase_user"]
    name = request.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Server name is required.")
    await ensure_profile(client, user)
    invite_code = generate_invite_code()
    server_id = str(uuid.uuid4())
    await client.rest(
        "POST",
        "servers",
        json_body={"id": server_id, "name": name[:80], "owner_id": user["id"], "invite_code": invite_code},
        prefer="return=minimal",
    )
    await client.rest(
        "POST",
        "server_members",
        json_body={"server_id": server_id, "user_id": user["id"], "role": "owner"},
        prefer="return=minimal",
    )
    channels = [{"server_id": server_id, "name": ch, "type": "text"} for ch in DEFAULT_CHANNELS]
    await client.rest("POST", "channels", json_body=channels, prefer="return=minimal")
    new_server = await get_server(client, server_id)
    _audit("create_server", server_id, user["id"], success=True)
    return {"success": True, "server": new_server}


@app.post("/api/servers/join")
async def join_server(request: JoinServerRequest, user=Depends(get_current_user)):
    client = user["supabase_user"]
    admin = supabase_admin()
    clean_code = request.invite_code.strip().upper()
    server_rows = await admin.rest("GET", "servers", params={"invite_code": f"eq.{clean_code}", "select": "*"})
    if not server_rows:
        raise HTTPException(status_code=404, detail="Invalid invite code. Server not found.")
    server = server_rows[0]

    bans = await admin.rest(
        "GET",
        "server_bans",
        params={"server_id": f"eq.{server['id']}", "user_id": f"eq.{user['id']}", "select": "id"},
    )
    if bans:
        raise HTTPException(status_code=403, detail="You are banned from this server.")

    existing = await get_server_member(client, server["id"], user["id"])
    if existing:
        return {"success": True, "server": server, "message": "You are already a member of this server!"}

    await ensure_profile(client, user)
    await client.rest(
        "POST",
        "server_members",
        json_body={"server_id": server["id"], "user_id": user["id"], "role": "member"},
        prefer="return=minimal",
    )
    _audit("join_server", server["id"], user["id"], success=True)
    return {"success": True, "server": server}


@app.post("/api/servers/{server_id}/channels")
async def create_channel(server_id: str, request: CreateChannelRequest, user=Depends(get_current_user)):
    client = user["supabase_user"]
    await require_server_permission(client, server_id, user["id"], "manage_channels")
    channel_type = request.type.strip().lower()
    if channel_type not in {"text", "voice"}:
        raise HTTPException(status_code=400, detail="Channel type must be text or voice.")
    channel = (await client.rest(
        "POST",
        "channels",
        json_body={"server_id": server_id, "name": normalize_channel_name(request.name), "type": channel_type},
        prefer="return=representation",
    ))[0]
    _audit("create_channel", server_id, user["id"], success=True)
    return {"success": True, "channel": channel}


@app.patch("/api/servers/{server_id}/members/{target_user_id}/role")
async def update_member_role(server_id: str, target_user_id: str, request: RoleUpdateRequest, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "manage_roles")
    target = await get_server_member(client, server_id, target_user_id)
    if not target:
        raise HTTPException(status_code=404, detail="Target user is not a server member.")
    old_role = target["role"]
    validate_role_change(user["id"], target_user_id, actor_role, old_role, request.role)
    updated = (await client.rest(
        "PATCH",
        "server_members",
        params={"server_id": f"eq.{server_id}", "user_id": f"eq.{target_user_id}"},
        json_body={"role": request.role},
        prefer="return=representation",
    ))[0]
    _audit("update_member_role", server_id, user["id"], target_user_id, old_role, request.role, True)
    return {"success": True, "member": updated}


@app.post("/api/servers/{server_id}/transfer-ownership")
async def transfer_ownership(server_id: str, request: TransferOwnershipRequest, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "transfer_ownership")
    if actor_role != "owner":
        raise HTTPException(status_code=403, detail="Only the server owner can transfer ownership.")
    if request.new_owner_id == user["id"]:
        raise HTTPException(status_code=400, detail="You already own this server.")
    target = await get_server_member(client, server_id, request.new_owner_id)
    if not target:
        raise HTTPException(status_code=400, detail="New owner must be a current server member.")
    await supabase_admin().rpc("transfer_server_ownership", {
        "p_server_id": server_id,
        "p_current_owner_id": user["id"],
        "p_new_owner_id": request.new_owner_id,
    })
    _audit("transfer_ownership", server_id, user["id"], request.new_owner_id, "owner", "owner", True)
    return {"success": True, "message": "Ownership transferred."}


@app.post("/api/servers/{server_id}/members/{target_user_id}/kick")
async def kick_member(server_id: str, target_user_id: str, request: ReasonRequest | None = None, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "kick_members")
    if target_user_id == user["id"]:
        raise HTTPException(status_code=400, detail="You cannot kick yourself.")
    target = await get_server_member(client, server_id, target_user_id)
    if not target:
        raise HTTPException(status_code=404, detail="Target user is not a server member.")
    if not can_manage_target_member(actor_role, target["role"], "kick"):
        raise HTTPException(status_code=403, detail="You cannot kick this member.")
    await cleanup_voice_presence(supabase_admin(), server_id, target_user_id)
    await client.rest(
        "DELETE",
        "server_members",
        params={"server_id": f"eq.{server_id}", "user_id": f"eq.{target_user_id}"},
    )
    _audit("kick_member", server_id, user["id"], target_user_id, target["role"], None, True)
    return {"success": True, "message": "Member kicked."}


@app.post("/api/servers/{server_id}/members/{target_user_id}/ban")
async def ban_member(server_id: str, target_user_id: str, request: ReasonRequest | None = None, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "ban_members")
    if target_user_id == user["id"]:
        raise HTTPException(status_code=400, detail="You cannot ban yourself.")
    target = await get_server_member(client, server_id, target_user_id)
    if target and not can_manage_target_member(actor_role, target["role"], "ban"):
        raise HTTPException(status_code=403, detail="You cannot ban this member.")
    server = await get_server(client, server_id)
    if target_user_id == server["owner_id"]:
        raise HTTPException(status_code=403, detail="The server owner cannot be banned.")

    existing_ban = await supabase_admin().rest(
        "GET",
        "server_bans",
        params={"server_id": f"eq.{server_id}", "user_id": f"eq.{target_user_id}", "select": "id"},
    )
    if existing_ban:
        raise HTTPException(status_code=409, detail="This user is already banned from the server.")

    await client.rest(
        "POST",
        "server_bans",
        json_body={
            "server_id": server_id,
            "user_id": target_user_id,
            "banned_by": user["id"],
            "reason": (request.reason if request else None),
        },
        prefer="return=minimal",
    )
    await cleanup_voice_presence(supabase_admin(), server_id, target_user_id)
    await client.rest(
        "DELETE",
        "server_members",
        params={"server_id": f"eq.{server_id}", "user_id": f"eq.{target_user_id}"},
    )
    _audit("ban_member", server_id, user["id"], target_user_id, target["role"] if target else None, None, True)
    return {"success": True, "message": "Member banned."}


@app.get("/api/servers/{server_id}/bans")
async def list_bans(server_id: str, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "ban_members")
    if actor_role != "owner":
        raise HTTPException(status_code=403, detail="Only the server owner can view bans.")
    bans = await client.rest("GET", "server_bans", params={"server_id": f"eq.{server_id}", "select": "*", "order": "created_at.desc"})
    user_ids = sorted({b["user_id"] for b in bans} | {b["banned_by"] for b in bans})
    profiles = {}
    if user_ids:
        profile_rows = await client.rest("GET", "profiles", params={"id": f"in.({','.join(user_ids)})", "select": "id,username,full_name,avatar_url,email"})
        profiles = {p["id"]: p for p in profile_rows}
    return {
        "bans": [
            {
                **ban,
                "profile": profiles.get(ban["user_id"]),
                "banned_by_profile": profiles.get(ban["banned_by"]),
            }
            for ban in bans
        ]
    }


@app.delete("/api/servers/{server_id}/bans/{target_user_id}")
async def unban_member(server_id: str, target_user_id: str, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "ban_members")
    if actor_role != "owner":
        raise HTTPException(status_code=403, detail="Only the server owner can unban users.")
    await client.rest("DELETE", "server_bans", params={"server_id": f"eq.{server_id}", "user_id": f"eq.{target_user_id}"})
    _audit("unban_member", server_id, user["id"], target_user_id, None, None, True)
    return {"success": True, "message": "User unbanned."}


@app.patch("/api/servers/{server_id}")
async def update_server(server_id: str, request: ServerUpdateRequest, user=Depends(get_current_user)):
    client = user["supabase_user"]
    await require_server_permission(client, server_id, user["id"], "manage_server")
    payload = {}
    if request.name is not None:
        name = request.name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="Server name is required.")
        payload["name"] = name[:80]
    if request.description is not None:
        payload["description"] = request.description.strip()[:500] if request.description else None
    if not payload:
        raise HTTPException(status_code=400, detail="No server settings were provided.")
    server = (await client.rest(
        "PATCH",
        "servers",
        params={"id": f"eq.{server_id}"},
        json_body=payload,
        prefer="return=representation",
    ))[0]
    _audit("update_server", server_id, user["id"], success=True)
    return {"success": True, "server": server}


@app.post("/api/servers/{server_id}/regenerate-invite")
async def regenerate_invite(server_id: str, user=Depends(get_current_user)):
    client = user["supabase_user"]
    await require_server_permission(client, server_id, user["id"], "manage_invites")
    invite_code = generate_invite_code()
    server = (await supabase_admin().rest(
        "PATCH",
        "servers",
        params={"id": f"eq.{server_id}"},
        json_body={"invite_code": invite_code},
        prefer="return=representation",
    ))[0]
    _audit("regenerate_invite", server_id, user["id"], success=True)
    return {"success": True, "server": server}


@app.delete("/api/servers/{server_id}")
async def delete_server(server_id: str, user=Depends(get_current_user)):
    client = user["supabase_user"]
    actor_role = await require_server_permission(client, server_id, user["id"], "delete_server")
    if actor_role != "owner":
        raise HTTPException(status_code=403, detail="Only the server owner can delete this server.")
    await client.rest("DELETE", "servers", params={"id": f"eq.{server_id}"})
    _audit("delete_server", server_id, user["id"], success=True)
    return {"success": True, "message": "Server deleted."}


# =====================================================================
# CACHING
# =====================================================================

# In-memory caches: keyed by (doc_id, cache_type[, question_hash])
_generation_cache: dict[str, dict] = {}


def _cache_key(doc_id: str, cache_type: str, extra: str = "") -> str:
    """Generate a deterministic cache key."""
    raw = f"{doc_id}:{cache_type}:{extra}"
    return hashlib.sha256(raw.encode()).hexdigest()


def _get_cached(doc_id: str, cache_type: str, extra: str = ""):
    key = _cache_key(doc_id, cache_type, extra)
    return _generation_cache.get(key)


def _set_cached(doc_id: str, cache_type: str, value, extra: str = ""):
    key = _cache_key(doc_id, cache_type, extra)
    _generation_cache[key] = value


# =====================================================================
# TURN CREDENTIALS (unchanged)
# =====================================================================

@app.get("/api/turn-credentials")
async def get_turn_credentials():
    if not METERED_SECRET_KEY:
        print("[TURN-API] METERED_SECRET_KEY environment variable is not set.")
        raise HTTPException(
            status_code=500,
            detail="METERED_SECRET_KEY environment variable is not set on the backend."
        )

    post_url = f"https://{METERED_DOMAIN}/api/v1/turn/credential"
    post_params = {"secretKey": METERED_SECRET_KEY}
    post_data = {"label": "studycord-dev", "expiryInSeconds": 3600}

    masked_secret = f"{METERED_SECRET_KEY[:4]}...{METERED_SECRET_KEY[-4:]}" if METERED_SECRET_KEY and len(METERED_SECRET_KEY) > 8 else "***"
    print(f"[TURN-API] Creating TURN credential. URL: {post_url}?secretKey={masked_secret}, Body: {post_data}")

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            # Step 1: Create TURN credential
            post_response = await client.post(post_url, params=post_params, json=post_data)
            post_response.raise_for_status()
            credential_info = post_response.json()
            
            api_key = credential_info.get("apiKey")
            if not api_key:
                print(f"[TURN-API] POST response did not contain 'apiKey'. Response: {credential_info}")
                raise HTTPException(
                    status_code=502,
                    detail="Failed to retrieve apiKey from Metered credential creation."
                )

            # Step 2: Fetch ICE servers using the api key
            get_url = f"https://{METERED_DOMAIN}/api/v1/turn/credentials"
            get_params = {"apiKey": api_key}
            
            masked_api_key = f"{api_key[:4]}...{api_key[-4:]}" if len(api_key) > 8 else "***"
            print(f"[TURN-API] Fetching ICE servers. URL: {get_url}?apiKey={masked_api_key}")

            get_response = await client.get(get_url, params=get_params)
            get_response.raise_for_status()
            
            ice_servers = get_response.json()
            return ice_servers
    except httpx.HTTPStatusError as e:
        print(f"[TURN-API] Metered API returned error status {e.response.status_code}: {e.response.text}")
        raise HTTPException(
            status_code=502,
            detail=f"Failed to fetch credentials from Metered API: {e.response.text}"
        )
    except Exception as e:
        print(f"[TURN-API] Connection error fetching Metered credentials: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal connection error: {str(e)}"
        )

# =====================================================================
# STUDYCORD BASIC RAG MVP ENDPOINTS
# =====================================================================

# In-memory store: doc_id -> { "db": FAISS, "text": str, "filename": str }
vector_stores = {}

async def extract_text_from_file(file: UploadFile) -> str:
    filename = file.filename
    content = await file.read()
    ext = filename.split(".")[-1].lower()
    
    if ext == "txt":
        return content.decode("utf-8", errors="ignore")
    
    elif ext == "pdf":
        from pypdf import PdfReader
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
        
    elif ext == "docx":
        import docx
        docx_file = io.BytesIO(content)
        doc = docx.Document(docx_file)
        text = ""
        for p in doc.paragraphs:
            if p.text:
                text += p.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        return text
    
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def parse_json_from_response(content: str):
    # Try direct parse
    try:
        return json.loads(content.strip())
    except json.JSONDecodeError:
        pass
        
    # Try to find json block using regex
    match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", content)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            pass
            
    # Try finding the first '[' or '{' and last ']' or '}'
    start_idx = min(content.find('{') if '{' in content else len(content), content.find('[') if '[' in content else len(content))
    end_idx = max(content.rfind('}') if '}' in content else -1, content.rfind(']') if ']' in content else -1)
    if start_idx < end_idx:
        try:
            return json.loads(content[start_idx:end_idx+1])
        except json.JSONDecodeError:
            pass
            
    raise ValueError("Response could not be parsed as valid JSON.")

class ChatRequest(BaseModel):
    question: str
    mode: str = "chat"
    doc_id: str

class DocRequest(BaseModel):
    doc_id: str

@app.post("/api/rag/upload")
async def rag_upload(file: UploadFile = File(...)):
    if not os.getenv("GOOGLE_API_KEY"):
        raise HTTPException(
            status_code=400,
            detail="GEMINI_API_KEY is not configured on the backend. Please add it to Backend/.env"
        )
        
    try:
        text = await extract_text_from_file(file)
        if not text.strip():
            raise HTTPException(status_code=400, detail="The uploaded file contains no readable text.")
            
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)
        
        # Convert to Documents
        documents = [Document(page_content=chunk, metadata={"source": file.filename}) for chunk in chunks]
        
        # Embeddings and FAISS — still using Gemini embeddings
        embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
        db = FAISS.from_documents(documents, embeddings)
        
        doc_id = str(uuid.uuid4())
        vector_stores[doc_id] = {
            "db": db,
            "text": text,
            "filename": file.filename
        }
        
        return {
            "status": "success",
            "doc_id": doc_id,
            "filename": file.filename,
            "num_chunks": len(chunks)
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"[RAG-UPLOAD] Error processing document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")


@app.post("/api/rag/chat")
async def rag_chat(request: ChatRequest):
    endpoint = "RAG-CHAT"
    print(f"[{endpoint}] Request received — model={OPENROUTER_MODEL}, question='{request.question[:80]}...'")
    start_time = time.time()

    doc_data = vector_stores.get(request.doc_id)
    if not doc_data:
        raise HTTPException(status_code=404, detail="Document not found or expired. Please upload it again.")

    # Check cache for repeated chat questions
    cached = _get_cached(request.doc_id, "chat", request.question)
    if cached is not None:
        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — CACHE HIT")
        return cached

    try:
        db = doc_data["db"]
        docs = db.similarity_search(request.question, k=4)
        context = "\n\n".join([doc.page_content for doc in docs])
        
        chat = get_rag_chat_model(temperature=0.3)
        prompt = (
            f"You are an AI Study Assistant for StudyCord. Answer the student's question based strictly on the provided document context. "
            f"If the answer cannot be found in the context, you may use your general knowledge to answer, but state clearly that it is not explicitly mentioned in the document.\n\n"
            f"Document Context:\n{context}\n\n"
            f"Student Question: {request.question}\n\n"
            f"Helpful Answer:"
        )
        
        response = await chat.ainvoke(prompt)
        result = {"answer": response.content}

        # Cache the result
        _set_cached(request.doc_id, "chat", result, request.question)

        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — GENERATED (model={OPENROUTER_MODEL})")
        return result
    except HTTPException:
        raise
    except Exception as e:
        print(f"[{endpoint}] Error: {str(e)}")
        raise _handle_openrouter_error(e, endpoint)


@app.post("/api/rag/summary")
async def rag_summary(request: DocRequest):
    endpoint = "RAG-SUMMARY"
    print(f"[{endpoint}] Request received — model={OPENROUTER_MODEL}, doc_id={request.doc_id[:12]}...")
    start_time = time.time()

    doc_data = vector_stores.get(request.doc_id)
    if not doc_data:
        raise HTTPException(status_code=404, detail="Document not found.")

    # Check cache
    cached = _get_cached(request.doc_id, "summary")
    if cached is not None:
        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — CACHE HIT")
        return cached

    try:
        text = doc_data["text"][:50000]  # truncate to stay within safe prompt length
        chat = get_rag_chat_model(temperature=0.2)
        
        prompt = (
            "Analyze the provided document text and generate a structured summary. "
            "You MUST reply ONLY with a valid JSON object matching the following structure:\n"
            "{\n"
            "  \"executive_summary\": \"A concise 2-3 paragraph summary of the entire document.\",\n"
            "  \"key_concepts\": [\n"
            "    {\"concept\": \"Name of concept\", \"description\": \"Definition/explanation of the concept\"}\n"
            "  ],\n"
            "  \"key_points\": [\n"
            "    \"Important point or takeaway 1\",\n"
            "    \"Important point or takeaway 2\"\n"
            "  ]\n"
            "}\n\n"
            "Do not include any formatting, markdown wrappers, or explanation outside of the raw JSON object.\n\n"
            f"Document Text:\n{text}"
        )
        
        response = await chat.ainvoke(prompt)
        parsed_json = parse_json_from_response(response.content)

        # Cache the result
        _set_cached(request.doc_id, "summary", parsed_json)

        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — GENERATED (model={OPENROUTER_MODEL})")
        return parsed_json
    except HTTPException:
        raise
    except ValueError as e:
        print(f"[{endpoint}] JSON parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail="The AI model returned a response that could not be parsed. Please try again.")
    except Exception as e:
        print(f"[{endpoint}] Error: {str(e)}")
        raise _handle_openrouter_error(e, endpoint)


@app.post("/api/rag/flashcards")
async def rag_flashcards(request: DocRequest):
    endpoint = "RAG-FLASHCARDS"
    print(f"[{endpoint}] Request received — model={OPENROUTER_MODEL}, doc_id={request.doc_id[:12]}...")
    start_time = time.time()

    doc_data = vector_stores.get(request.doc_id)
    if not doc_data:
        raise HTTPException(status_code=404, detail="Document not found.")

    # Check cache
    cached = _get_cached(request.doc_id, "flashcards")
    if cached is not None:
        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — CACHE HIT")
        return cached

    try:
        text = doc_data["text"][:50000]
        chat = get_rag_chat_model(temperature=0.3)
        
        prompt = (
            "Analyze the provided document text and generate a list of 5 to 8 high-quality revision flashcards. "
            "You MUST reply ONLY with a valid JSON array matching the following structure:\n"
            "[\n"
            "  {\"question\": \"A clear, specific study question?\", \"answer\": \"A concise, informative answer explaining the concept.\"}\n"
            "]\n\n"
            "Do not include any formatting, markdown wrappers, or explanation outside of the raw JSON array.\n\n"
            f"Document Text:\n{text}"
        )
        
        response = await chat.ainvoke(prompt)
        parsed_json = parse_json_from_response(response.content)
        result = {"flashcards": parsed_json}

        # Cache the result
        _set_cached(request.doc_id, "flashcards", result)

        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — GENERATED (model={OPENROUTER_MODEL})")
        return result
    except HTTPException:
        raise
    except ValueError as e:
        print(f"[{endpoint}] JSON parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail="The AI model returned a response that could not be parsed. Please try again.")
    except Exception as e:
        print(f"[{endpoint}] Error: {str(e)}")
        raise _handle_openrouter_error(e, endpoint)


@app.post("/api/rag/mcq")
async def rag_mcq(request: DocRequest):
    endpoint = "RAG-MCQ"
    print(f"[{endpoint}] Request received — model={OPENROUTER_MODEL}, doc_id={request.doc_id[:12]}...")
    start_time = time.time()

    doc_data = vector_stores.get(request.doc_id)
    if not doc_data:
        raise HTTPException(status_code=404, detail="Document not found.")

    # Check cache
    cached = _get_cached(request.doc_id, "mcq")
    if cached is not None:
        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — CACHE HIT")
        return cached

    try:
        text = doc_data["text"][:50000]
        chat = get_rag_chat_model(temperature=0.3)
        
        prompt = (
            "Analyze the provided document text and generate 5 multiple-choice questions (MCQs) for revision. "
            "Each question must have exactly 4 unique options, and one correct answer (which must exactly match one of the options). "
            "You MUST reply ONLY with a valid JSON array matching the following structure:\n"
            "[\n"
            "  {\n"
            "    \"question\": \"Question text here?\",\n"
            "    \"options\": [\"Option 1\", \"Option 2\", \"Option 3\", \"Option 4\"],\n"
            "    \"correct_answer\": \"Option 2\"\n"
            "  }\n"
            "]\n\n"
            "Do not include any formatting, markdown wrappers, or explanation outside of the raw JSON array.\n\n"
            f"Document Text:\n{text}"
        )
        
        response = await chat.ainvoke(prompt)
        parsed_json = parse_json_from_response(response.content)
        result = {"mcqs": parsed_json}

        # Cache the result
        _set_cached(request.doc_id, "mcq", result)

        elapsed = time.time() - start_time
        print(f"[{endpoint}] Completed in {elapsed:.2f}s — GENERATED (model={OPENROUTER_MODEL})")
        return result
    except HTTPException:
        raise
    except ValueError as e:
        print(f"[{endpoint}] JSON parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail="The AI model returned a response that could not be parsed. Please try again.")
    except Exception as e:
        print(f"[{endpoint}] Error: {str(e)}")
        raise _handle_openrouter_error(e, endpoint)

if __name__ == "__main__":
    # pyrefly: ignore [missing-import]
    import uvicorn
    # Read port from env if needed, default to 8000
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
