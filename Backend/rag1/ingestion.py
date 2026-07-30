"""Authenticated local ingestion for personal RAG 1 documents."""

import io
import re
import time
import uuid
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from fastapi import UploadFile
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .db import RagDocument, RagDocumentRepository
from .paths import (
    DOCUMENT_TEXT_FILENAME,
    FAISS_DOCSTORE_FILENAME,
    FAISS_INDEX_FILENAME,
    create_staging_directory,
    promote_staging_directory,
    remove_document_directory,
    remove_staging_directory,
    validate_uuid,
)


MAX_UPLOAD_BYTES = 20 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
ARTIFACT_VERSION = 1
EMBEDDING_MODEL = "models/gemini-embedding-001"

SUPPORTED_EXTENSIONS = frozenset({"pdf", "txt", "docx"})
ALLOWED_MIME_TYPES = {
    "pdf": frozenset({"application/pdf", "application/octet-stream"}),
    "txt": frozenset({"text/plain", "application/octet-stream"}),
    "docx": frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/zip",
            "application/octet-stream",
        }
    ),
}


class RagIngestionError(Exception):
    def __init__(self, status_code: int, detail: str):
        super().__init__(detail)
        self.status_code = status_code
        self.detail = detail


@dataclass(frozen=True)
class ValidatedUpload:
    display_filename: str
    detected_type: str
    content: bytes


@dataclass(frozen=True)
class IngestionResult:
    doc_id: str
    filename: str
    detected_type: str
    size_bytes: int
    chunk_count: int
    text: str
    vector_store: FAISS


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_display_filename(filename: str | None) -> str:
    if not filename:
        raise RagIngestionError(400, "The uploaded file must have a filename.")
    basename = re.split(r"[\\/]", filename)[-1]
    cleaned = "".join(
        character
        for character in basename
        if character.isprintable() and character not in "\x00\r\n"
    ).strip()
    if not cleaned:
        raise RagIngestionError(400, "The uploaded file must have a valid filename.")
    return cleaned[:255]


def _validate_upload_content(
    filename: str,
    content: bytes,
    *,
    supplied_mime: str = "",
    expected_type: str | None = None,
) -> ValidatedUpload:
    display_filename = _safe_display_filename(filename)
    extension = (
        display_filename.rsplit(".", 1)[-1].lower()
        if "." in display_filename
        else ""
    )
    if extension not in SUPPORTED_EXTENSIONS:
        raise RagIngestionError(
            400,
            "Unsupported file format. Upload a PDF, TXT, or DOCX document.",
        )

    if not content:
        raise RagIngestionError(400, "The uploaded file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise RagIngestionError(413, "Study Helper documents must be 20 MB or smaller.")

    canonical_expected_type = (expected_type or "").strip().lower()
    if canonical_expected_type and canonical_expected_type != extension:
        raise RagIngestionError(
            400,
            "The document type does not match its filename.",
        )

    canonical_mime = supplied_mime.strip().lower()
    if canonical_mime and canonical_mime not in ALLOWED_MIME_TYPES[extension]:
        raise RagIngestionError(
            400,
            "The uploaded file type does not match a supported document format.",
        )

    if extension == "pdf":
        if not content.startswith(b"%PDF-"):
            raise RagIngestionError(400, "The uploaded PDF is not valid.")
    elif extension == "docx":
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                if (
                    "[Content_Types].xml" not in names
                    or "word/document.xml" not in names
                ):
                    raise RagIngestionError(400, "The uploaded DOCX is not valid.")
                if sum(info.file_size for info in archive.infolist()) > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise RagIngestionError(400, "The uploaded DOCX expands beyond the safe limit.")
                if archive.testzip() is not None:
                    raise RagIngestionError(400, "The uploaded DOCX is corrupt.")
        except (zipfile.BadZipFile, OSError) as error:
            raise RagIngestionError(400, "The uploaded DOCX is not valid.") from error
    elif b"\x00" in content:
        raise RagIngestionError(400, "The uploaded TXT file does not contain plain text.")

    return ValidatedUpload(
        display_filename=display_filename,
        detected_type=extension,
        content=content,
    )


async def _read_and_validate_upload(file: UploadFile) -> ValidatedUpload:
    content = await file.read(MAX_UPLOAD_BYTES + 1)
    return _validate_upload_content(
        file.filename or "",
        content,
        supplied_mime=file.content_type or "",
    )


def _extract_text(upload: ValidatedUpload) -> str:
    try:
        if upload.detected_type == "txt":
            text = upload.content.decode("utf-8", errors="ignore")
        elif upload.detected_type == "pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(upload.content))
            text = "\n".join(
                page_text
                for page in reader.pages
                if (page_text := page.extract_text())
            )
        else:
            import docx

            document = docx.Document(io.BytesIO(upload.content))
            sections: list[str] = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text
            ]
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        sections.append(" | ".join(row_text))
            text = "\n".join(sections)
    except RagIngestionError:
        raise
    except Exception as error:
        raise RagIngestionError(
            400,
            f"The uploaded {upload.detected_type.upper()} document is corrupt or unreadable.",
        ) from error

    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise RagIngestionError(
            400,
            "The uploaded file contains no readable text.",
        )
    return normalized


def _build_embeddings():
    return GoogleGenerativeAIEmbeddings(model=EMBEDDING_MODEL)


def _build_vector_store(documents: list[Document], embeddings) -> FAISS:
    return FAISS.from_documents(documents, embeddings)


def _validate_staged_artifacts(staging_directory: Path) -> None:
    for filename in (
        DOCUMENT_TEXT_FILENAME,
        FAISS_INDEX_FILENAME,
        FAISS_DOCSTORE_FILENAME,
    ):
        artifact = staging_directory / filename
        if (
            not artifact.is_file()
            or artifact.is_symlink()
            or artifact.stat().st_size <= 0
        ):
            raise RagIngestionError(
                500,
                "Document processing did not produce complete local artifacts.",
            )


async def _ingest_validated_upload(
    upload: ValidatedUpload,
    user_id: str,
    *,
    repository: RagDocumentRepository | None = None,
) -> IngestionResult:
    """Validate, process, and atomically persist one authenticated RAG 1 document."""
    started_at = time.time()
    canonical_user_id = str(validate_uuid(user_id, "user id"))
    doc_id = str(uuid.uuid4())
    repository = repository or RagDocumentRepository()
    staging_directory: Path | None = None
    promoted = False
    metadata_created = False
    completed = False
    stage = "validation"

    try:
        created_at = _utc_now()
        stage = "metadata"
        repository.create(
            RagDocument(
                id=doc_id,
                user_id=canonical_user_id,
                original_filename=upload.display_filename,
                detected_type=upload.detected_type,
                size_bytes=len(upload.content),
                status="processing",
                chunk_count=None,
                artifact_version=ARTIFACT_VERSION,
                created_at=created_at,
                updated_at=created_at,
            )
        )
        metadata_created = True

        stage = "extraction"
        text = _extract_text(upload)
        staging_directory = create_staging_directory(doc_id, repository.data_dir)
        (staging_directory / DOCUMENT_TEXT_FILENAME).write_text(
            text,
            encoding="utf-8",
        )

        stage = "chunking"
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
        )
        chunks = [chunk for chunk in splitter.split_text(text) if chunk.strip()]
        if not chunks:
            raise RagIngestionError(
                400,
                "The uploaded file produced no usable text chunks.",
            )
        documents = [
            Document(
                page_content=chunk,
                metadata={"source": upload.display_filename},
            )
            for chunk in chunks
        ]

        stage = "embedding"
        embeddings = _build_embeddings()
        vector_store = _build_vector_store(documents, embeddings)

        stage = "faiss_save"
        vector_store.save_local(str(staging_directory))
        _validate_staged_artifacts(staging_directory)

        stage = "promotion"
        promote_staging_directory(
            staging_directory,
            canonical_user_id,
            doc_id,
            repository.data_dir,
        )
        staging_directory = None
        promoted = True

        stage = "ready"
        if not repository.update_status_for_user(
            doc_id,
            canonical_user_id,
            "ready",
            _utc_now(),
            chunk_count=len(chunks),
        ):
            raise RagIngestionError(
                500,
                "Document metadata could not be finalized.",
            )

        completed = True
        elapsed = time.time() - started_at
        print(
            f"[RAG-INGEST] doc_id={doc_id} status=ready "
            f"chunks={len(chunks)} elapsed={elapsed:.2f}s"
        )
        return IngestionResult(
            doc_id=doc_id,
            filename=upload.display_filename,
            detected_type=upload.detected_type,
            size_bytes=len(upload.content),
            chunk_count=len(chunks),
            text=text,
            vector_store=vector_store,
        )
    except RagIngestionError:
        raise
    except Exception as error:
        if stage == "embedding":
            raise RagIngestionError(
                502,
                "Document embeddings could not be generated.",
            ) from error
        raise RagIngestionError(
            500,
            "The document could not be processed.",
        ) from error
    finally:
        if metadata_created and not completed:
            if staging_directory is not None:
                try:
                    remove_staging_directory(staging_directory, repository.data_dir)
                except Exception as cleanup_error:
                    print(
                        f"[RAG-INGEST] doc_id={doc_id} stage={stage} "
                        f"staging_cleanup=failed error={type(cleanup_error).__name__}"
                    )
            if promoted:
                try:
                    remove_document_directory(
                        canonical_user_id,
                        doc_id,
                        repository.data_dir,
                    )
                except Exception as cleanup_error:
                    print(
                        f"[RAG-INGEST] doc_id={doc_id} stage={stage} "
                        f"artifact_cleanup=failed error={type(cleanup_error).__name__}"
                    )
            try:
                repository.delete_for_user(doc_id, canonical_user_id)
            except Exception as cleanup_error:
                print(
                    f"[RAG-INGEST] doc_id={doc_id} stage={stage} "
                    f"metadata_cleanup=failed error={type(cleanup_error).__name__}"
                )


async def ingest_rag_document_bytes(
    content: bytes,
    filename: str,
    detected_type: str,
    user_id: str,
    *,
    repository: RagDocumentRepository | None = None,
) -> IngestionResult:
    """Ingest trusted backend-owned bytes through the normal RAG 1 pipeline."""
    upload = _validate_upload_content(
        filename,
        content,
        expected_type=detected_type,
    )
    return await _ingest_validated_upload(
        upload,
        user_id,
        repository=repository,
    )


async def ingest_rag_document(
    file: UploadFile,
    user_id: str,
    *,
    repository: RagDocumentRepository | None = None,
) -> IngestionResult:
    """Adapt a browser UploadFile to the shared RAG 1 ingestion core."""
    try:
        upload = await _read_and_validate_upload(file)
        return await _ingest_validated_upload(
            upload,
            user_id,
            repository=repository,
        )
    finally:
        try:
            await file.close()
        except Exception:
            pass
