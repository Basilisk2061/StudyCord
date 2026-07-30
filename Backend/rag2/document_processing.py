"""Content-authoritative validation, extraction, and chunking for RAG 2."""

import io
import zipfile
from dataclasses import dataclass

from langchain_text_splitters import RecursiveCharacterTextSplitter


MAX_RESOURCE_BYTES = 10 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_CHUNKS = 2_000
MAX_STORED_CHUNK_CHARS = 10_000
CHUNK_SIZE = 1_000
CHUNK_OVERLAP = 200
SUPPORTED_TYPES = frozenset({"pdf", "docx", "txt"})


class Rag2DocumentError(Exception):
    def __init__(
        self,
        status_code: int,
        detail: str,
        *,
        detected_type: str | None = None,
    ):
        super().__init__(detail)
        self.status_code = status_code
        self.detail = detail
        self.detected_type = detected_type


@dataclass(frozen=True)
class ProcessedDocument:
    detected_type: str
    text: str
    chunks: list[str]


def _candidate_type(filename: str) -> str:
    clean_name = filename.strip()
    extension = clean_name.rsplit(".", 1)[-1].lower() if "." in clean_name else ""
    if extension not in SUPPORTED_TYPES:
        raise Rag2DocumentError(
            422,
            "The resource filename is not a supported RAG 2 candidate.",
        )
    return extension


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            try:
                if reader.decrypt("") == 0:
                    raise Rag2DocumentError(
                        422,
                        "Encrypted PDFs must be unlocked before indexing.",
                        detected_type="pdf",
                    )
            except Rag2DocumentError:
                raise
            except Exception as error:
                raise Rag2DocumentError(
                    422,
                    "Encrypted PDFs must be unlocked before indexing.",
                    detected_type="pdf",
                ) from error
        text = "\n".join(
            page_text
            for page in reader.pages
            if (page_text := page.extract_text())
        )
    except Rag2DocumentError:
        raise
    except Exception as error:
        raise Rag2DocumentError(
            422,
            "The PDF is corrupt or unreadable.",
            detected_type="pdf",
        ) from error
    return text


def _extract_docx(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = set(archive.namelist())
            if (
                "[Content_Types].xml" not in names
                or "word/document.xml" not in names
            ):
                raise Rag2DocumentError(
                    422,
                    "The DOCX does not contain the required Word document structure.",
                )
            if (
                sum(info.file_size for info in archive.infolist())
                > MAX_DOCX_UNCOMPRESSED_BYTES
            ):
                raise Rag2DocumentError(
                    422,
                    "The DOCX expands beyond the safe indexing limit.",
                    detected_type="docx",
                )
            if archive.testzip() is not None:
                raise Rag2DocumentError(
                    422,
                    "The DOCX is corrupt.",
                    detected_type="docx",
                )
    except Rag2DocumentError:
        raise
    except (zipfile.BadZipFile, OSError) as error:
        raise Rag2DocumentError(422, "The DOCX is corrupt or unreadable.") from error

    try:
        import docx

        document = docx.Document(io.BytesIO(content))
        sections = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text
        ]
        for table in document.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    sections.append(" | ".join(row_text))
        return "\n".join(sections)
    except Exception as error:
        raise Rag2DocumentError(
            422,
            "The DOCX is corrupt or unreadable.",
            detected_type="docx",
        ) from error


def _extract_txt(content: bytes) -> str:
    if b"\x00" in content:
        raise Rag2DocumentError(422, "The TXT resource contains binary data.")
    try:
        return content.decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError as error:
        raise Rag2DocumentError(
            422,
            "The TXT resource is not valid UTF-8 text.",
        ) from error


def _detect_and_extract(content: bytes) -> tuple[str, str]:
    if b"%PDF-" in content[:1024]:
        return "pdf", _extract_pdf(content)

    if zipfile.is_zipfile(io.BytesIO(content)):
        return "docx", _extract_docx(content)

    return "txt", _extract_txt(content)


def process_resource_document(filename: str, content: bytes) -> ProcessedDocument:
    """Validate actual bytes, extract normalized text, and split deterministically."""
    candidate_type = _candidate_type(filename)
    if not content:
        raise Rag2DocumentError(422, "The Storage object is empty.")
    if len(content) > MAX_RESOURCE_BYTES:
        raise Rag2DocumentError(
            413,
            "RAG 2 resources must be 10 MB or smaller.",
        )

    detected_type, extracted_text = _detect_and_extract(content)
    if detected_type != candidate_type:
        raise Rag2DocumentError(
            422,
            "The resource content does not match its candidate filename extension.",
            detected_type=detected_type,
        )

    normalized = extracted_text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise Rag2DocumentError(
            422,
            "The resource contains no readable text.",
            detected_type=detected_type,
        )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = [chunk for chunk in splitter.split_text(normalized) if chunk.strip()]
    if not chunks:
        raise Rag2DocumentError(
            422,
            "The resource produced no usable text chunks.",
            detected_type=detected_type,
        )
    if len(chunks) > MAX_CHUNKS:
        raise Rag2DocumentError(
            413,
            "The resource produces too many chunks for indexing.",
            detected_type=detected_type,
        )
    if any(len(chunk) > MAX_STORED_CHUNK_CHARS for chunk in chunks):
        raise Rag2DocumentError(
            413,
            "A resource chunk exceeds the defensive storage limit.",
            detected_type=detected_type,
        )

    return ProcessedDocument(
        detected_type=detected_type,
        text=normalized,
        chunks=chunks,
    )
